"""Create a polished Word version of the 2026 FFHub interview preparation guide."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "FFHub_Interview_Prep_2026-08-04.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "20252B"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "CAD3DF"
WHITE = "FFFFFF"
GOLD = "7A5A00"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text, size=11, color=INK, italic=False):
    """Add simple Markdown-style bold spans to a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size=6, left_color=None, left_size=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(left_size if edge == "start" and left_size else size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), left_color if edge == "start" and left_color else color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(
                cell,
                CELL_MARGIN_TOP_BOTTOM,
                CELL_MARGIN_START_END,
                CELL_MARGIN_TOP_BOTTOM,
                CELL_MARGIN_START_END,
            )


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    # OOXML requires abstract numbering definitions before concrete w:num nodes.
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, size=11, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=11, color=INK)
    else:
        add_rich_text(p, text)
    return p


def add_body(doc, text, bold_lead=None, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, size=11, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=11, color=color, italic=italic)
    else:
        add_rich_text(p, text, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_font(run, 16, BLUE, True)
    elif level == 2:
        set_font(run, 13, BLUE, True)
    else:
        set_font(run, 12, DARK_BLUE, True)
    return p


def add_callout(doc, label, text, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_GRAY)
    set_cell_borders(cell, color=BORDER, size=4, left_color=accent, left_size=18)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.18
    label_run = p.add_run(label.upper() + "  ")
    set_font(label_run, size=9, color=accent, bold=True)
    text_run = p.add_run(text)
    set_font(text_run, size=11, color=INK, italic=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_row_cant_split(table.rows[0])
    items = [
        ("FORMAT", "Group interview"),
        ("DURATION", "30-45 minutes"),
        ("TIME", "8pm UK / 9pm CEST"),
        ("PLATFORM", "Riverside FM"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, items):
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_borders(cell, color=BORDER, size=4)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        label_run = p.add_run(label + "\n")
        set_font(label_run, size=8, color=BLUE, bold=True)
        value_run = p.add_run(value)
        set_font(value_run, size=9.5, color=NAVY, bold=True)
    return table


def add_quick_table(doc):
    rows = [
        ("Best overall asset", "Kimi Antonelli", "429 points; 20+ in 10 of 11"),
        ("Safest premium", "Lewis Hamilton", "No negative scores or retirements"),
        ("Biggest improver", "Isack Hadjar", "1.2 first-five avg to 22.6 last-five"),
        ("Premium watch", "Lando Norris", "115 points across the last three"),
        ("Fantasy-only gem", "Esteban Ocon", "120 fantasy points; only 3 F1 points"),
        ("Thank-you-and-sell", "Franco Colapinto", "+£5.0M gained; weak forward value"),
        ("Price disappointment", "Oscar Piastri", "£24.4M; 160 points; 3 retirements"),
        ("Cheap constructor watch", "Audi", "0.6 first-five avg to 16.4 last-five"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ("Question", "Answer", "One-line evidence")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, 9.5, NAVY, True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for category, answer, evidence in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for cell, text in zip(cells, (category, answer, evidence)):
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, 9.4, INK, bold=(text == answer))
    set_table_geometry(table, [2350, 2200, 4810])
    return table


def add_sources_table(doc):
    rows = [
        ("Fantasy totals", "data/seed/official_fantasy_points.json"),
        ("Scoring components", "web/public/data/actual_round*.json"),
        ("Prices", "data/seed/fantasy_prices.json (after round 13)"),
        ("Zandvoort", "web/public/data/predictions_round14_pre_fp.json"),
        ("R15-R19 priors", "web/public/data/horizon_projections.json"),
        ("Budget valuation", "data/experiments/budget_point_value_2026.md"),
        ("F1 standings", "formula1.com/en/results/2026/drivers"),
        ("Official 2026 calendar", "formula1.com/en/racing/2026"),
        ("Malaysia confirmation", "formula1.com - F1/FIA confirmation, 26 July 2026"),
        ("Imola contingency", "Sky Sport / The Race - 25 July 2026 reports"),
    ]
    table = doc.add_table(rows=1, cols=2)
    for cell, text in zip(table.rows[0].cells, ("Evidence", "Source")):
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_borders(cell)
        run = cell.paragraphs[0].add_run(text)
        set_font(run, 9.5, NAVY, True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for label, source in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for index, text in enumerate((label, source)):
            set_cell_borders(cells[index])
            run = cells[index].paragraphs[0].add_run(text)
            set_font(run, 9.2, INK, bold=(index == 0))
    set_table_geometry(table, [2500, 6860])
    return table


def add_calendar_budget_table(doc):
    rows = [
        ("Old working horizon", "11", "1.39 pts", "5.34 pts"),
        ("Confirmed calendar with Malaysia", "12", "1.39 pts", "5.35 pts"),
        ("True extra race on top", "13", "1.39 pts", "5.35 pts"),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = (
        "Calendar case",
        "Races left",
        "Forecast +£0.3M",
        "Secured +£1.0M",
    )
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_borders(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_font(run, 9.2, NAVY, True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for label, races, forecast_value, secured_value in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for index, text in enumerate((label, races, forecast_value, secured_value)):
            set_cell_borders(cells[index])
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(text)
            set_font(run, 9.2, INK, bold=(index == 0))
    set_table_geometry(table, [3300, 1500, 2280, 2280])
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    for header in (section.header, section.even_page_header, section.first_page_header):
        hp = header.paragraphs[0]
        for existing_run in list(hp.runs):
            hp._p.remove(existing_run._r)
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)

    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        fp = footer.paragraphs[0]
        for existing_run in list(fp.runs):
            fp._p.remove(existing_run._r)
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_after = Pt(0)
        label = fp.add_run("4 August 2026   |   Page ")
        set_font(label, 9, MUTED)
        add_page_number(fp)


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    doc.core_properties.title = "FFHub Interview Prep - BoxBox F1 Fantasy"
    doc.core_properties.subject = "2026 F1 Fantasy guest interview preparation"
    doc.core_properties.author = "BoxBox F1 Fantasy"
    doc.core_properties.keywords = "F1 Fantasy, interview, 2026, BoxBox"

    bullet_id = add_numbering_definition(doc, "bullet")
    number_id = add_numbering_definition(doc, "decimal")

    # Workshop-agenda first-page pattern, resolved through compact_reference_guide.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("GUEST INTERVIEW GUIDE")
    set_font(run, 10, BLUE, True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("FFHub Interview Prep")
    set_font(run, 28, NAVY, True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("2026 season insights, asset watchlist and camera-ready answers")
    set_font(run, 13.2, MUTED)
    add_metric_strip(doc)
    add_body(
        doc,
        "Prepared for Quintin (BoxBox F1 Fantasy), appearing with Euan (F1 Fantasy Pole Position) and Luke (F1 Fantasy Strategy).",
        italic=True,
        color=MUTED,
        after=8,
    )
    add_callout(
        doc,
        "Best soundbite",
        "The edge is separating signal from story: use repeatable FP pace, price reliability properly, and sell the asset that made you money once its forward value disappears.",
    )

    add_heading(doc, "Five numbers worth remembering", 1)
    five_numbers = [
        "**Kimi Antonelli: 429 fantasy points.** He is 85 ahead of Hamilton, has 20+ in 10 of 11 weekends, and has no retirement.",
        "**Lewis Hamilton: the safest premium.** No negative weekend, no retirement, ten 20-point rounds, and the lowest volatility among the leading seven drivers.",
        "**Isack Hadjar: the biggest turnaround.** His average rose from 1.2 across his first five completed races to 22.6 across his last five.",
        "**Sprint weekends: 70% more total driver points.** The best-driver ceiling has also been 25% higher than on normal weekends.",
        "**A forecast +£0.3M rise is worth about 1.4 future points.** Budget is option value, not points already banked.",
    ]
    for item in five_numbers:
        p = doc.add_paragraph()
        apply_numbering(p, number_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_rich_text(p, item)

    add_heading(doc, "Opening and season-so-far answers", 1)
    add_heading(doc, "Suggested introduction", 2)
    add_callout(
        doc,
        "On camera",
        "BoxBox is my attempt to turn F1 Fantasy from pure gut feel into a proper decision tool. We combine historical form, circuit characteristics, practice pace, weather, reliability and Monte Carlo simulation. A model cannot remove uncertainty, but it can price it and show when the popular story does not match the numbers.",
    )
    add_body(doc, "Add the number of seasons you have personally played and your social channels.", italic=True, color=MUTED)

    add_heading(doc, "Strategy and biggest lesson", 2)
    add_callout(
        doc,
        "On camera",
        "The biggest lesson is that budget only matters when it crosses an affordability threshold. I do not want to give away ten real points just to chase a routine £0.3M rise. Build budget early, but later in the season convert it into points.",
    )
    add_body(
        doc,
        "Supporting stat: two simulated managers finished with the same £124.3M team value but 150 points apart. Having spending power and deploying it well are different skills.",
    )

    add_heading(doc, "Chip success and failure", 2)
    add_body(doc, "Use only the example that matches your real team.", italic=True, color=MUTED)
    chip_points = [
        "**3x benchmark:** Kimi's 68 in China was the best driver round. A 3x there added 68 points over the normal 2x captain score; his 62 in Canada was second.",
        "**Monaco Limitless lesson:** perfect-hindsight cap removal was worth only 18 points, the smallest uplift all year. Miami was worth 167. Premium track did not mean premium assets won.",
        "**No Negative hindsight:** the largest all-asset negative damage came in Australia (146), Barcelona (135) and Monaco (128). Trigger it from incident, weather and reliability risk - not simply because it is a Sprint.",
    ]
    for item in chip_points:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Reliability: the season's hidden currency", 2)
    add_body(
        doc,
        "There have been 48 DNF/DSQ/DNS results across 242 car-starts: 4.36 per race, or roughly 19.8% attrition. Kimi and Hamilton have zero. Norris and Piastri have three each, costing each 60 fantasy points in retirement penalties alone.",
    )

    doc.add_page_break()
    add_heading(doc, "Asset analysis and watchlist", 1)
    add_quick_table(doc)

    add_heading(doc, "Asset that impressed: Hamilton", 2)
    add_callout(
        doc,
        "On camera",
        "Kimi is obviously the asset of the season, but Hamilton has impressed me most as a fantasy asset. No negative score, no retirement and ten 20-point weekends out of eleven. In a year of reliability shocks, his floor has been as valuable as anyone's ceiling.",
    )
    add_body(
        doc,
        "Constructor alternative: Racing Bulls has 344 points and rose £6.6M from £6.3M to £12.9M. Even now, it has returned 26.7 season points per current £M - close to Mercedes and Ferrari.",
    )

    add_heading(doc, "Premium watch: Norris", 2)
    add_callout(
        doc,
        "On camera",
        "Norris is my premium watchlist name. His first-five average was 19.0, his last-five average is 29.8, and he has 115 across the last three. Piastri scored 27 over those same three. Norris is still £1.1M cheaper than at the start, and our six-round model has him second only to Kimi.",
    )

    add_heading(doc, "Mid-price watch: Hadjar - with caution", 2)
    add_body(
        doc,
        "Hadjar's +21.4 points-per-race improvement is the largest in the field. His recent three-race average is also 4.7 higher than Verstappen's. But the forward prior is nearer 9.7 per race for R14-R19. Watch the FP evidence; do not blindly extrapolate the streak.",
    )

    add_heading(doc, "Fantasy-only gem: Ocon", 2)
    add_body(
        doc,
        "Ocon is 17th in the real championship on three points but 12th in fantasy on 120. He has never retired or posted a negative fantasy weekend, and 70 of his points came from overtakes - second only to Verstappen's 75. It is the clearest example of why F1 standings and fantasy value are different.",
    )

    add_heading(doc, "Thank-you-and-sell candidate: Colapinto", 2)
    add_body(
        doc,
        "Colapinto delivered 123 points, no negative weekend, no retirement and the largest driver price gain at +£5.0M. But he now costs £11.2M and the R14-R19 model returns only 3.2 total points. He was a superb budget builder; the forward value may already be gone.",
    )

    add_heading(doc, "Current-price disappointment: Piastri", 2)
    add_body(
        doc,
        "Piastri has 160 points, three retirements, three negative weekends and only three 20-point rounds. At £24.4M, he is near Hamilton/Leclerc money, yet the six-round prior gives him 55.8 points versus Hamilton's 145.4, Leclerc's 134.2 and Norris's 167.2.",
    )

    add_heading(doc, "Constructor watch: Audi", 2)
    add_body(
        doc,
        "Audi averaged 0.6 points over its first five completed races and 16.4 over its last five. At £6.8M, the R14-R19 model gives it 57.8 points - the best current low-cost forward watch. Racing Bulls remains the season's value success, but the priors expect it to fall after Zandvoort.",
    )

    add_heading(doc, "Pivotal moments", 1)
    pivotal = [
        "**China:** Kimi's 68 was the highest driver score and the cleanest 3x benchmark.",
        "**Miami:** Norris scored 54 and perfect-hindsight Limitless value peaked at 167.",
        "**Monaco:** only 122 points across all 22 drivers; ten assets were negative; Limitless cap value was only 18.",
        "**Barcelona:** Hamilton scored 56 while the field absorbed 135 points of negative damage.",
        "**Silverstone:** 384 driver points and only 23 points of negative damage - a positive-multiplier weekend rather than a protection weekend.",
        "**The Norris/Hadjar turn:** Norris has 115 over the last three; Hadjar has 76 and has outscored Verstappen's 62.",
    ]
    for item in pivotal:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Second-half opportunities", 1)
    add_heading(doc, "The two remaining Sprint weekends", 2)
    add_body(
        doc,
        "Completed Sprint weekends generated 70% more total driver points and a 25% higher best-driver ceiling. This favours positive multiplier chips, but it does not automatically make a Sprint the best No Negative or Limitless round.",
    )
    add_callout(
        doc,
        "Current prior",
        "Singapore is the larger multiplier opportunity: Kimi projects at 50.6 there versus 30.5 at Zandvoort. The Singapore number is priors-only, so use Zandvoort FP1 evidence or retain the chip rather than choosing from calendar type alone.",
        GOLD,
    )

    add_heading(doc, "Madrid debut", 2)
    add_body(
        doc,
        "BoxBox currently rates Madrid as balanced: overtaking 5/10, average corner speed 6/10, straight-line importance 6/10 and downforce 6/10, with a higher 7/10 Turn 1 incident risk.",
    )
    add_callout(
        doc,
        "On camera",
        "For a debut circuit I want optionality. Hold transfers, avoid committing a chip on track reputation, and let FP supply the first comparable evidence. Repeatable three- and five-lap pace matters more than one purple lap.",
    )

    add_heading(doc, "Malaysia / Sepang", 2)
    add_body(
        doc,
        "F1 and the FIA have confirmed that Sepang will host the 2026 Bahrain Grand Prix on 2-4 October, between Baku and Singapore. The official calendar now contains 23 rounds, so there are 12 races remaining after Hungary rather than the 11 used in the original budget study. It also creates a demanding Baku-Sepang-Singapore sequence.",
    )
    add_callout(
        doc,
        "Important caveat",
        "BoxBox's current seed and horizon do not yet include Sepang. Do not quote a Malaysia projection as model output.",
        RED,
    )

    add_heading(doc, "Calendar changes and the points-to-budget exchange rate", 2)
    add_body(
        doc,
        "The marginal-budget study fits a saturating curve: 8.57 × (1 - exp(-races remaining / 1.80)) future projected points per secured £1M, followed by separate discounts for forecast price realisation and realised marginal points. A forecast rise becomes spendable one race later.",
    )
    add_calendar_budget_table(doc)
    add_body(
        doc,
        "The headline is that Malaysia does not materially increase the present exchange rate. At this long horizon the curve is already almost flat: moving from 11 to 12 races raises the decision value of a secured £1M by only about 0.01 point, while a forecast +£0.3M still rounds to 1.39 points. A forecast rise still needs to be roughly £2.16M to justify giving up 10 certain points under the smooth model.",
    )
    add_callout(
        doc,
        "On camera",
        "Malaysia gives us another scoring and chip opportunity, but it does not suddenly make routine price rises more valuable. At 12 races left, a forecast +£0.3M is still worth only about 1.4 future points. The extra date matters more for chip timing and transfer flexibility than for the budget exchange rate.",
        GOLD,
    )

    add_heading(doc, "How to frame the Imola rumour", 3)
    add_body(
        doc,
        "Imola is being discussed as a contingency finale if the Middle East ending changes. Current reporting does not establish it as a 24th race: it could replace Abu Dhabi one-for-one, or become one replacement if both Qatar and Abu Dhabi are lost. In the first case the horizon stays at 12 races; in the second it falls back to 11. Only a genuine additional Imola round on top of the current 23 would create a 13-race horizon - and even then the current forecast +£0.3M value still rounds to 1.39 points.",
    )
    add_body(
        doc,
        "The calendar effect becomes more important later. With three races left, a forecast +£0.3M is worth about 0.93 point; adding one genuinely usable race lifts that to about 1.13. Recalculate once the finale is confirmed rather than treating today's 12-race assumption as permanent.",
    )

    add_heading(doc, "Where is the edge?", 2)
    add_callout(
        doc,
        "On camera",
        "Use repeatable FP pace rather than the fastest single lap, model DNF risk instead of calling every retirement random, and price budget as an option rather than as points. Then sell the asset that made you money when its forward value disappears.",
    )

    add_heading(doc, "Bold predictions", 1)
    add_heading(doc, "F1 prediction", 2)
    add_callout(
        doc,
        "Bold call",
        "Lando Norris wins at least two more Grands Prix and finishes third in the Drivers' Championship. He is fifth now, 32 points behind Russell, but Hungary looked like the start of a genuine second-half charge.",
    )
    add_heading(doc, "F1 Fantasy prediction", 2)
    add_callout(
        doc,
        "Bold call",
        "Norris is the number-two fantasy driver over the remainder of the season, behind Kimi but ahead of Hamilton, Leclerc and Verstappen. Ferrari becomes the best premium-constructor value rather than Mercedes.",
    )
    add_body(
        doc,
        "Current R14-R19 priors: Kimi 221.1, Norris 167.2, Hamilton 145.4, Verstappen 137.1 and Leclerc 134.2. Ferrari projects 284.5 at £26.6M versus Mercedes 343.3 at £32.6M - almost identical projected points per £M.",
    )

    add_heading(doc, "Rapid-fire answers", 1)
    rapid = [
        ("Most consistent premium", "Hamilton"),
        ("Best overall asset", "Kimi"),
        ("Best improvement", "Hadjar"),
        ("Best premium rebound", "Norris"),
        ("Best fantasy-not-F1 player", "Ocon"),
        ("Best budget-building success", "Colapinto / Racing Bulls"),
        ("Current-price disappointment", "Piastri"),
        ("Cheap constructor watch", "Audi"),
        ("Biggest chip lesson", "Limitless needs expensive assets to dominate"),
        ("One strategic rule", "Do not spend 10 points chasing a routine £0.3M rise"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for question, answer in rapid:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for index, text in enumerate((question, answer)):
            set_cell_borders(cells[index])
            if index == 0:
                set_cell_fill(cells[index], LIGHT_BLUE)
            run = cells[index].paragraphs[0].add_run(text)
            set_font(run, 9.6, NAVY if index == 0 else INK, bold=(index == 0))
    set_table_geometry(table, [3150, 6210])
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=30, start=120, bottom=30, end=120)

    doc.add_page_break()
    add_heading(doc, "Sources and caveats", 1)
    add_sources_table(doc)
    caveats = [
        "Perfect-hindsight chip numbers select the best legal £100M or uncapped lineup from realised scores. They measure opportunity, not forecast skill.",
        "R14-R19 values are model priors, not promises, and the current horizon file does not yet include Sepang.",
        "The interview outline's 'first 15 races' should be phrased as 11 completed Grands Prix. The confirmed Malaysia date makes 12 remaining on the official 23-round calendar.",
        "Imola is a reported contingency, not a confirmed extra race. Recalculate the budget exchange rate when the final calendar is settled.",
        "Personal chip success/failure examples should be used only when they match the manager's actual team history.",
    ]
    for item in caveats:
        add_list_item(doc, item, bullet_id)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
